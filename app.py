import streamlit as st
import google.generativeai as genai
from docx import Document
from docx.shared import Pt
from io import BytesIO
from supabase import create_client, Client

# --- CONFIGURACIÓN DE APIS ---
try:
    genai.configure(api_key=st.secrets["GENAI_KEY"])
    model = genai.GenerativeModel('gemini-3-flash-preview')
    supabase: Client = create_client(st.secrets["SUPABASE_URL"], st.secrets["SUPABASE_KEY"])
except Exception as e:
    st.error("Error de configuración. Revisa tus Secrets en Streamlit.")

# --- CONFIGURACIÓN DE PÁGINA ---
st.set_page_config(page_title="Asistente Planeador Pro", page_icon="🍎", layout="wide")

# Ocultar menú de Streamlit
st.markdown("""
    <style>
    /* Oculta el header pero permite que la flecha lateral siga funcionando */
    [data-testid="stHeader"] {background: rgba(0,0,0,0); height: 0rem;}
    #MainMenu {visibility: hidden;}
    footer {visibility: hidden;}
    /* Ajuste para que el contenido no quede pegado arriba */
    .block-container {padding-top: 2rem;}
    </style>
    """, unsafe_allow_html=True)

# --- LÓGICA DE USUARIO (SUPABASE AUTH) ---
if 'user' not in st.session_state:
    st.session_state.user = None

def registrar_usuario(email, password):
    try:
        supabase.auth.sign_up({"email": email, "password": password})
        st.success("¡Cuenta creada! Ya puedes iniciar sesión.")
    except Exception as e:
        st.error(f"Error al registrar: {e}")

def iniciar_sesion(email, password):
    try:
        res = supabase.auth.sign_in_with_password({"email": email, "password": password})
        st.session_state.user = res.user
        st.rerun()
    except Exception as e:
        st.error("Credenciales incorrectas o usuario no encontrado.")

def cerrar_sesion():
    supabase.auth.sign_out()
    st.session_state.user = None
    st.session_state.resultado = None # Limpiamos pantalla al salir
    st.rerun()

# --- FUNCIONES DE BASE DE DATOS ---
def guardar_en_db(email, tema, grado, metodologia, contenido):
    data = {
        "user_email": email, 
        "tema": tema, 
        "grado": grado, 
        "metodologia": metodologia, 
        "contenido": contenido
    }
    supabase.table("planeaciones").insert(data).execute()

def obtener_historial(email):
    res = supabase.table("planeaciones").select("*").eq("user_email", email).order("created_at", desc=True).execute()
    return res.data

def eliminar_de_db(id_registro):
    try:
        supabase.table("planeaciones").delete().eq("id", id_registro).execute()
        st.toast("Consulta eliminada correctamente", icon="🗑️")
        st.rerun() # Esto recarga la app para actualizar la lista
    except Exception as e:
        st.error(f"No se pudo eliminar: {e}")

# --- FUNCIÓN WORD PROFESIONAL ---
def crear_word(titulo, contenido):
    doc = Document()
    doc.add_heading(f"Planeación: {titulo}", 0)
    for linea in contenido.split('\n'):
        linea_limpia = linea.replace('**', '').strip()
        if not linea_limpia: continue
        
        if linea.strip().startswith('#'):
            p = doc.add_paragraph()
            run = p.add_run(linea_limpia.replace('#', ''))
            run.bold = True
            run.font.size = Pt(14)
        else:
            doc.add_paragraph(linea_limpia)
    bio = BytesIO()
    doc.save(bio)
    return bio.getvalue()

# --- BARRA LATERAL (AUTH Y HISTORIAL) ---
with st.sidebar:
    st.image("https://cdn-icons-png.flaticon.com/512/3429/3429153.png", width=80)
    if st.session_state.user is None:
        st.header("🔐 Acceso")
        tab1, tab2 = st.tabs(["Ingresar", "Registrarse"])
        with tab1:
            e_log = st.text_input("Correo", key="l_email")
            p_log = st.text_input("Contraseña", type="password", key="l_pass")
            if st.button("Entrar", use_container_width=True): iniciar_sesion(e_log, p_log)
        with tab2:
            e_reg = st.text_input("Correo", key="r_email")
            p_reg = st.text_input("Contraseña", type="password", key="r_pass")
            if st.button("Crear Cuenta", use_container_width=True): registrar_usuario(e_reg, p_reg)
    else:
        st.write(f"Docente: **{st.session_state.user.email}**")
        if st.button("Cerrar Sesión"): cerrar_sesion()
        st.divider()
        st.subheader("📚 Mis Planeaciones")
        historial = obtener_historial(st.session_state.user.email)
        
        if historial:
            for p in historial:
                # Usamos el tema y la fecha para identificar la consulta
                with st.expander(f"📄 {p['tema']}"):
                    st.caption(f"Grado: {p['grado']} | {p['metodologia']}")
                    
                    col_ver, col_del = st.columns(2)
                    
                    # Botón para CARGAR la consulta
                    if col_ver.button("👁️ Ver", key=f"v_{p['id']}", use_container_width=True):
                        st.session_state.resultado = p['contenido']
                        st.session_state.tema_guardado = p['tema']
                    
                    # Botón para ELIMINAR la consulta (con color rojo)
                    if col_del.button("🗑️", key=f"d_{p['id']}", use_container_width=True, help="Eliminar de forma permanente"):
                        eliminar_de_db(p['id'])
        else:
            st.info("Sin registros.")
        

# --- SECCIÓN DE HERRAMIENTAS EXTERNAS ---
with st.sidebar:
    st.divider()
    st.subheader("🛠️ Herramientas de Apoyo")
    st.caption("Usa estas herramientas para complementar tu planeación con material visual y juegos:")

    # Botón para Sopas de Letras y Crucigramas
    st.link_button(
        "🧩 Crear Sopa de Letras", 
        "https://www.educima.com/wordsearch.php", 
        use_container_width=True,
        help="Copia las palabras de tu planeación y pégalas aquí para generar el juego."
    )

    # Botón para generar imágenes/dibujos para colorear
    st.link_button(
        "🎨 Dibujos para Colorear (IA)", 
        "https://www.bing.com/images/create", 
        use_container_width=True,
        help="Pide a la IA: 'Un dibujo lineal para colorear sobre [tu tema] para niños'."
    )

    # Botón para dinámicas interactivas
    st.link_button(
        "🎲 Juegos Interactivos", 
        "https://wordwall.net/es", 
        use_container_width=True,
        help="Crea cuestionarios, ruedas de la fortuna y juegos digitales rápido."
    )

# --- CUERPO PRINCIPAL ---
st.title("🍎 Asistente de Planeaciones Pro 📚")

if st.session_state.user is None:
    st.info("👈 Por favor, regístrate o inicia sesión para comenzar a guardar tus planeaciones.")
else:
    # 1. DATOS DE LA CLASE (Recuperado)
    with st.container(border=True):
        st.subheader("1. Configuración")
        col1, col2 = st.columns(2)
        with col1:
            tema = st.text_input("¿Qué tema vas a enseñar?", placeholder="Ej. Fracciones equivalentes")
        with col2:
            grado = st.selectbox("¿Para qué grado?", ["Preescolar", "1º Primaria", "2º Primaria", "3º Primaria", "4º Primaria", "5º Primaria", "6º Primaria", "Secundaria"])
        
        metodologias = ["Constructivismo", "Nueva escuela Mexicana", "Montessori", "Humanista", "Aprendizaje Basado en Problemas", "Enfoque por Competencias", "Aula Invertida", "Sistema Preventivo (Salesianas)"]
        metodologia = st.selectbox("Metodología", metodologias)

    # 2. OPCIONES EXTRAS (Recuperado)
    with st.expander("🛠️ 2. Opciones Adicionales"):
        col_opt1, col_opt2, col_opt3 = st.columns(3)
        incluir_quiz = col_opt1.checkbox("Añadir un Quiz rápido")
        incluir_examen = col_opt2.checkbox("Preguntas de examen")
        incluir_proyecto = col_opt3.checkbox("Proyecto mensual")

    if st.button("Generar Planeación ✨", use_container_width=True, type="primary"):
        if tema:
            with st.spinner('Diseñando la clase mágica...'):
                instrucciones_extra = ""
                if incluir_quiz: instrucciones_extra += " - Incluye un quiz rápido de 5 preguntas.\n"
                if incluir_examen: instrucciones_extra += " - Sugiere 5 preguntas de examen con respuestas.\n"
                if incluir_proyecto: instrucciones_extra += " - Propón una idea para un proyecto mensual.\n"

                # TU PROMPT COMPLETO
                prompt = (
                            f"Eres un experto en diseño curricular. Genera una secuencia didáctica estructurada sobre: '{tema}' para {grado}. "
                            f"Es FUNDAMENTAL que toda la planeación respete los principios de la metodología: {metodologia}.\n\n"
                            f"INSTRUCCIONES DE ESTRUCTURA:\n"
                            f"1. Si '{tema}' incluye varios conceptos o es un proceso complejo (ej. 'La exposición'), divide la planeación en una secuencia lógica de pasos o sesiones (ej. Sesión 1: Introducción/Concepto, Sesión 2: Preparación, Sesión 3: Práctica/Aplicación).\n"
                            f"2. Para cada sesión o paso, incluye: un Objetivo específico, y Actividades claras con su tiempo estimado.\n"
                            f"3. Proporciona un Resumen general al inicio de todo el documento.\n"
                            f"4. Añade actividades complementarias (ideas para hojas lúdicas o de trabajo).\n"
                            f"5. Incluye fuentes de apoyo (libros, citas) y sugiere links reales o términos de búsqueda precisos para videos de YouTube.\n"
                            f"6. Responde 100% en español.\n\n"
                            f"EXTRAS SOLICITADOS:\n{instrucciones_extra}\n\n"
                            f"REGLA DE FORMATO: Sé directo. Comienza inmediatamente con el título de la planeación. NO saludes, NO des introducciones conversacionales y NO confirmes tu rol como experto."
                        )

                
                
                response = model.generate_content(prompt)
                st.session_state.resultado = response.text
                st.session_state.tema_guardado = tema
                
                # GUARDADO AUTOMÁTICO
                guardar_en_db(st.session_state.user.email, tema, grado, metodologia, response.text)
                st.toast("¡Planeación guardada!", icon="✅")
        else:
            st.warning("Escribe un tema primero.")

    # 3. MOSTRAR RESULTADOS
    if 'resultado' in st.session_state and st.session_state.resultado:
        st.divider()
        with st.container(border=True):
            st.markdown(st.session_state.resultado)
            archivo_word = crear_word(st.session_state.tema_guardado, st.session_state.resultado)
            st.download_button(label="📥 Descargar en Word", data=archivo_word, file_name=f"Planeacion_{st.session_state.tema_guardado}.docx")